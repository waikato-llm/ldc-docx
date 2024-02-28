import argparse
from typing import Iterable, List, Union

import docx

from wai.logging import LOGGING_WARNING
from seppl.io import locate_files
from ldc.core import domain_suffix
from ldc.api.pretrain import PretrainData, PretrainReader


class DOCXPretrainReader(PretrainReader):
    """
    Extracts text from MS Word .docx files to use for pretraining.
    """

    def __init__(self, source: Union[str, List[str]] = None, source_list: Union[str, List[str]] = None,
                 output_paragraphs: bool = False,
                 logger_name: str = None, logging_level: str = LOGGING_WARNING):
        """
        Initializes the reader.

        :param source: the filename(s)
        :param source_list: the file(s) with filename(s)
        :param output_paragraphs: whether to output paragraphs rather than the full document
        :type output_paragraphs: bool
        :param logger_name: the name to use for the logger
        :type logger_name: str
        :param logging_level: the logging level to use
        :type logging_level: str
        """
        super().__init__(logger_name=logger_name, logging_level=logging_level)
        self.source = source
        self.source_list = source_list
        self.output_paragraphs = output_paragraphs
        self._inputs = None
        self._current_input = None
        self._current_doc = None

    def name(self) -> str:
        """
        Returns the name of the reader, used as command-line name.

        :return: the name
        :rtype: str
        """
        return "from-docx-" + domain_suffix(self)

    def description(self) -> str:
        """
        Returns a description of the reader.

        :return: the description
        :rtype: str
        """
        return "Extracts text from MS Word .docx files to use for pretraining."

    def _create_argparser(self) -> argparse.ArgumentParser:
        """
        Creates an argument parser. Derived classes need to fill in the options.

        :return: the parser
        :rtype: argparse.ArgumentParser
        """
        parser = super()._create_argparser()
        parser.add_argument("-i", "--input", type=str, help="Path to the MS Word .docx file(s) to read; glob syntax is supported", required=False, nargs="*")
        parser.add_argument("-I", "--input_list", type=str, help="Path to the text file(s) listing the MS Word .docx files to use", required=False, nargs="*")
        parser.add_argument("--output_paragraphs", action="store_true", help="Whether to output individual paragraphs rather than whole documents.", required=False)
        return parser

    def _apply_args(self, ns: argparse.Namespace):
        """
        Initializes the object with the arguments of the parsed namespace.

        :param ns: the parsed arguments
        :type ns: argparse.Namespace
        """
        super()._apply_args(ns)
        self.source = ns.input
        self.source_list = ns.input_list
        self.output_paragraphs = ns.output_paragraphs

    def initialize(self):
        """
        Initializes the reading, e.g., for opening files or databases.
        """
        super().initialize()
        self._inputs = locate_files(self.source, input_lists=self.source_list, fail_if_empty=True)
        self._current_doc = None

    def read(self) -> Iterable[PretrainData]:
        """
        Loads the data and returns the items one by one.

        :return: the data
        :rtype: Iterable[PretrainData]
        """
        self.finalize()

        self._current_input = self._inputs.pop(0)
        self.session.current_input = self._current_input
        self.logger().info("Reading from: " + str(self.session.current_input))

        try:
            self._current_doc = docx.Document(self.session.current_input)

            if self.output_paragraphs:
                for i, para in enumerate(self._current_doc.paragraphs):
                    meta = dict()
                    meta["file"] = self.session.current_input
                    meta["paragraph"] = i
                    yield PretrainData(
                        content=para.text,
                        meta=meta,
                    )
            else:
                lines = [para.text for para in self._current_doc.paragraphs]
                meta = dict()
                meta["file"] = self.session.current_input
                yield PretrainData(
                    content="\n".join(lines),
                    meta=meta,
                )
            self._current_doc = None
        except:
            self.logger().exception("Failed to read from: %s" % self.session.current_input)
            self._current_doc = None
            yield None

    def has_finished(self) -> bool:
        """
        Returns whether reading has finished.

        :return: True if finished
        :rtype: bool
        """
        return (len(self._inputs) == 0) and (self._current_doc is None)

    def finalize(self):
        """
        Finishes the reading, e.g., for closing files or databases.
        """
        if self._current_input is not None:
            super().finalize()
            self._current_input = None
            self._current_doc = None
